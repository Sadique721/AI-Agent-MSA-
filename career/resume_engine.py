"""
career/resume_engine.py
========================
Resume & Cover Letter Intelligence Engine (V7).

Manages four resume variants per job:
  1. master_resume  — full canonical resume (source of truth)
  2. ats_resume     — keyword-optimised plain-text for ATS systems
  3. targeted       — company + role specific, generated via LLM
  4. cover_letter   — personalised cover letter, generated via LLM

All versions are persisted as versioned JSON in data/resumes/.
PDF/DOCX export available via python-docx (already in requirements.txt).

Usage:
    from career.resume_engine import ResumeEngine
    engine = ResumeEngine()
    engine.load_master("...full resume text...")
    cv = engine.generate_cover_letter(job_listing)
    targeted = engine.generate_targeted_resume(job_listing)
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
from datetime import datetime, timezone
from typing import Dict, List, Optional

from career.job_models import JobListing, ResumeVersion
from config import RESUME_DIR, USER_PROFILE

logger = logging.getLogger("msa.career.resume")


class ResumeEngine:
    """
    Central resume management system with version history
    and LLM-powered targeted generation.
    """

    def __init__(self, llm_manager=None) -> None:
        self._llm = llm_manager
        self._master_text: str = ""
        self._versions: Dict[str, ResumeVersion] = {}
        os.makedirs(RESUME_DIR, exist_ok=True)
        self._load_master_from_disk()

    # ── Master Resume ─────────────────────────────────────────────────────────

    def load_master(self, resume_text: str) -> None:
        """Set and persist the master resume."""
        self._master_text = resume_text.strip()
        version = self._save_version("master", resume_text)
        logger.info("[ResumeEngine] Master resume loaded (%d chars)", len(resume_text))

    def get_master(self) -> str:
        """Return the current master resume text."""
        return self._master_text

    # ── ATS Resume ────────────────────────────────────────────────────────────

    def generate_ats_resume(self, job: JobListing) -> str:
        """
        Generate a keyword-optimised ATS resume.
        Injects missing job keywords into the skills section of the master.
        Does NOT use LLM — purely additive keyword injection.
        """
        if not self._master_text:
            return ""

        from career.resume_matcher import ResumeMatcher
        matcher = ResumeMatcher()
        gaps = matcher.identify_gaps(job.description, self._master_text)

        ats_text = self._master_text
        if gaps:
            keyword_line = "Additional relevant skills: " + ", ".join(gaps[:15])
            # Inject after the SKILLS section header if found
            if "skills" in ats_text.lower():
                import re
                ats_text = re.sub(
                    r"(?i)(skills[:\s].*?\n)",
                    r"\1" + keyword_line + "\n",
                    ats_text,
                    count=1,
                )
            else:
                ats_text += f"\n\n{keyword_line}"

        version = self._save_version(f"ats_{job.id}", ats_text, job_id=job.id)
        logger.info("[ResumeEngine] ATS resume generated for job %s", job.id)
        return ats_text

    # ── Targeted Resume ───────────────────────────────────────────────────────

    def generate_targeted_resume(
        self, job: JobListing, llm_manager=None
    ) -> str:
        """
        LLM-powered company + role specific resume.
        Falls back to ATS resume if no LLM available.
        """
        llm = llm_manager or self._llm
        if llm is None or not self._master_text:
            return self.generate_ats_resume(job)

        prompt = (
            f"You are an expert resume writer. Rewrite the following resume to be "
            f"specifically tailored for this job posting:\n\n"
            f"JOB TITLE: {job.title}\n"
            f"COMPANY: {job.company}\n"
            f"JOB DESCRIPTION:\n{job.description[:1500]}\n\n"
            f"ORIGINAL RESUME:\n{self._master_text[:2000]}\n\n"
            f"Instructions:\n"
            f"- Keep all factual information intact (do NOT fabricate experience)\n"
            f"- Reorder bullet points to match job requirements\n"
            f"- Use keywords from the job description naturally\n"
            f"- Maintain professional tone\n"
            f"- Output ONLY the resume text, no explanations\n"
        )
        try:
            targeted = llm.generate(prompt, max_tokens=800)
            version = self._save_version(
                f"targeted_{job.company.replace(' ', '_')}_{job.id}",
                targeted,
                job_id=job.id,
            )
            logger.info("[ResumeEngine] Targeted resume generated for %s @ %s", job.title, job.company)
            return targeted
        except Exception as exc:
            logger.warning("[ResumeEngine] LLM targeted generation failed: %s", exc)
            return self.generate_ats_resume(job)

    # ── Cover Letter ──────────────────────────────────────────────────────────

    def generate_cover_letter(
        self, job: JobListing, llm_manager=None
    ) -> str:
        """
        Generate a personalised cover letter.
        Falls back to a structured template if LLM unavailable.
        """
        llm = llm_manager or self._llm
        name = USER_PROFILE.get("name", "Md Sadique Amin")
        role = USER_PROFILE.get("role", "Software Engineer")
        skills_list = ", ".join(USER_PROFILE.get("skills", [])[:8])

        if llm is None:
            cover_letter = self._template_cover_letter(job, name, role, skills_list)
        else:
            prompt = (
                f"Write a professional and personalised cover letter for the following job:\n\n"
                f"JOB TITLE: {job.title}\n"
                f"COMPANY: {job.company}\n"
                f"LOCATION: {job.location}\n"
                f"JOB DESCRIPTION:\n{job.description[:1000]}\n\n"
                f"APPLICANT:\n"
                f"  Name: {name}\n"
                f"  Role: {role}\n"
                f"  Key Skills: {skills_list}\n\n"
                f"Instructions:\n"
                f"- 3-4 paragraphs, professional but warm tone\n"
                f"- Highlight why this specific company / role excites the applicant\n"
                f"- Reference 2-3 specific skills from the job description\n"
                f"- End with a clear call-to-action\n"
                f"- Output ONLY the letter text\n"
            )
            try:
                cover_letter = llm.generate(prompt, max_tokens=500)
            except Exception as exc:
                logger.warning("[ResumeEngine] LLM cover letter failed: %s", exc)
                cover_letter = self._template_cover_letter(job, name, role, skills_list)

        version = self._save_version(
            f"cover_letter_{job.id}", cover_letter, job_id=job.id
        )
        return cover_letter

    @staticmethod
    def _template_cover_letter(job, name, role, skills) -> str:
        """Structured template fallback."""
        today = datetime.now().strftime("%B %d, %Y")
        return (
            f"{today}\n\n"
            f"Hiring Manager\n"
            f"{job.company}\n\n"
            f"Dear Hiring Manager,\n\n"
            f"I am writing to express my interest in the {job.title} position at {job.company}. "
            f"As a {role} with expertise in {skills}, I am excited about the opportunity to "
            f"contribute to your team.\n\n"
            f"My background aligns closely with the requirements outlined in the job description. "
            f"I have hands-on experience with many of the technologies and practices you are looking for, "
            f"and I am confident in my ability to deliver meaningful contributions from day one.\n\n"
            f"I would welcome the opportunity to discuss how my skills and experience can benefit "
            f"{job.company}. Thank you for considering my application.\n\n"
            f"Sincerely,\n{name}"
        )

    # ── Version History ───────────────────────────────────────────────────────

    def get_version_history(self) -> List[ResumeVersion]:
        """Return all saved versions sorted by creation date."""
        return sorted(self._versions.values(), key=lambda v: v.created_at or "", reverse=True)

    def get_version(self, version_id: str) -> Optional[ResumeVersion]:
        return self._versions.get(version_id)

    # ── Persistence ───────────────────────────────────────────────────────────

    def _save_version(
        self, label: str, content: str, job_id: Optional[str] = None
    ) -> ResumeVersion:
        version_id = hashlib.sha256(
            f"{label}_{datetime.now(timezone.utc).isoformat()}".encode()
        ).hexdigest()[:12]
        rv = ResumeVersion(
            version_id=version_id,
            label=label,
            content=content,
            job_id=job_id,
            created_at=datetime.now(timezone.utc).isoformat(),
        )
        self._versions[version_id] = rv
        # Persist to disk
        path = os.path.join(RESUME_DIR, f"{label}_{version_id}.json")
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(rv.to_dict(), f, indent=2)
        except Exception as exc:
            logger.warning("[ResumeEngine] Could not save version to disk: %s", exc)
        return rv

    def _load_master_from_disk(self) -> None:
        """Load the most recent 'master' resume from disk on startup."""
        try:
            candidates = [
                f for f in os.listdir(RESUME_DIR)
                if f.startswith("master_") and f.endswith(".json")
            ]
            if not candidates:
                return
            latest = sorted(candidates)[-1]
            with open(os.path.join(RESUME_DIR, latest), encoding="utf-8") as f:
                data = json.load(f)
            self._master_text = data.get("content", "")
            logger.info("[ResumeEngine] Loaded master resume from %s", latest)
        except Exception as exc:
            logger.debug("[ResumeEngine] No master resume on disk: %s", exc)

    def export_to_docx(self, version_id: str, output_path: str) -> bool:
        """Export a version to .docx format using python-docx."""
        version = self._versions.get(version_id)
        if not version:
            logger.warning("[ResumeEngine] Version %s not found", version_id)
            return False
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(USER_PROFILE.get("name", "Resume"), 0)
            for line in version.content.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())
            doc.save(output_path)
            logger.info("[ResumeEngine] Exported %s to %s", version_id, output_path)
            return True
        except ImportError:
            logger.warning("[ResumeEngine] python-docx not installed — skipping DOCX export")
            return False
        except Exception as exc:
            logger.error("[ResumeEngine] DOCX export failed: %s", exc)
            return False
