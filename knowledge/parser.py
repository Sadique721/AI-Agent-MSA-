"""
knowledge/parser.py
===================
Document parsers supporting TXT, PDF, DOCX, Markdown, JSON, Source Code, and GitHub repos.
Extracts clean plain text and returns structured documents.
"""

import os
import json
import logging
import zipfile
import tempfile
import shutil
import re
from typing import List, Dict, Any, Optional

logger = logging.getLogger("msa.knowledge.parser")


class DocumentParser:
    """
    Parses various file types into clean text ready for chunking.
    """

    @staticmethod
    def clean_text(text: str) -> str:
        """Standardize text spacing and remove junk characters."""
        if not text:
            return ""
        # Replace vertical tabs/form feeds, multiple spaces, multiple newlines
        text = re.sub(r"[\r\v\f]", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n+", "\n\n", text)
        return text.strip()

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a single file based on its extension.
        Returns a dict: { "text": str, "metadata": dict }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        base_name = os.path.basename(file_path)
        metadata = {
            "source": file_path,
            "filename": base_name,
            "file_size": os.path.getsize(file_path),
            "extension": ext
        }

        try:
            if ext == ".txt":
                text = self._parse_txt(file_path)
            elif ext == ".pdf":
                text = self._parse_pdf(file_path)
            elif ext in (".docx", ".doc"):
                text = self._parse_docx(file_path)
            elif ext == ".json":
                text, extra_meta = self._parse_json(file_path)
                metadata.update(extra_meta)
            elif ext in (".md", ".markdown"):
                text = self._parse_markdown(file_path)
            elif ext in (".py", ".java", ".js", ".ts", ".html", ".css", ".go", ".c", ".cpp", ".h", ".sh", ".ps1", ".properties", ".xml", ".yaml", ".yml"):
                text = self._parse_source_code(file_path)
            else:
                # Default text file parsing attempt
                text = self._parse_txt(file_path)
                metadata["warning"] = "unknown extension parsed as plain text"
            
            cleaned_text = self.clean_text(text)
            return {
                "text": cleaned_text,
                "metadata": metadata
            }
        except Exception as e:
            logger.error("Parser: failed to parse '%s' (%s)", file_path, e)
            raise RuntimeError(f"Parsing failed for {base_name}: {e}")

    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text files with encoding detection fallbacks."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("utf-8", b"", 0, 0, "Unable to decode text file")

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF using pypdf, falling back to basic extraction if needed."""
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages_text.append(t)
            return "\n\n".join(pages_text)
        except Exception as e:
            logger.warning("pypdf extraction failed for '%s' (%s). Trying basic binary fallback.", file_path, e)
            return self._parse_pdf_fallback(file_path)

    def _parse_pdf_fallback(self, file_path: str) -> str:
        """Crude regex-based pdf fallback if library fails."""
        # PDF is binary, but we can search for stream tags and try to find plain text
        try:
            with open(file_path, "rb") as f:
                content = f.read()
            # Find plaintext strings inside parentheses in PDF streams
            matches = re.findall(b"\\(([^)]+)\\)", content)
            text_parts = []
            for m in matches:
                try:
                    text_parts.append(m.decode("utf-8", errors="ignore"))
                except Exception:
                    pass
            return " ".join(text_parts)
        except Exception as e:
            logger.error("PDF fallback parsing failed: %s", e)
            return ""

    def _parse_docx(self, file_path: str) -> str:
        """Parse Word documents using docx (python-docx), fallback to zip XML parser."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            # Handle tables too
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.warning("python-docx failed for '%s' (%s). Trying XML zip fallback.", file_path, e)
            return self._parse_docx_fallback(file_path)

    def _parse_docx_fallback(self, file_path: str) -> str:
        """Docx is just a zip archive containing document.xml. Parse it directly."""
        try:
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read("word/document.xml").decode("utf-8")
                # Remove all XML tags to keep raw text
                clean_xml = re.sub(r"<[^>]+>", " ", xml_content)
                # Normalize spaces
                return clean_xml
        except Exception as e:
            logger.error("Docx zip fallback parsing failed: %s", e)
            return ""

    def _parse_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse JSON. Returns formatted text representation and extra metadata."""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # Pretty print representation
        text_rep = json.dumps(data, indent=2, ensure_ascii=False)
        
        extra_metadata = {}
        if isinstance(data, dict):
            # Extract simple metadata fields if present
            for k in ("title", "author", "description", "category", "tags"):
                if k in data:
                    extra_metadata[f"json_{k}"] = data[k]
        return text_rep, extra_metadata

    def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown as text."""
        return self._parse_txt(file_path)

    def _parse_source_code(self, file_path: str) -> str:
        """Parse source code with line number indicators / metadata decoration."""
        code_text = self._parse_txt(file_path)
        base_name = os.path.basename(file_path)
        # Wrap it with some semantic indicators to explain it's code
        return f"File: {base_name}\n```\n{code_text}\n```"


class GitHubRepositoryIndexer:
    """
    Downloads or clones a GitHub repository locally, parses all compatible files,
    and returns a list of Document objects ready for ingestion.
    """

    def __init__(self, parser: Optional[DocumentParser] = None):
        self.parser = parser or DocumentParser()

    def index_repo(self, repo_url: str, branch: str = "main") -> List[Dict[str, Any]]:
        """
        Download repo as ZIP via GitHub API, extract, and parse documents.
        Supports urls like: https://github.com/user/repo
        """
        import requests
        
        # Clean up URL
        repo_url = repo_url.strip().rstrip("/")
        if not repo_url.startswith("http"):
            # Assume user/repo format
            repo_url = f"https://github.com/{repo_url}"

        # Extract owner and repo name
        parts = repo_url.split("/")
        if len(parts) < 5:
            raise ValueError(f"Invalid GitHub repository URL: {repo_url}")
        
        owner, repo = parts[-2], parts[-1]
        zip_url = f"https://github.com/{owner}/{repo}/archive/refs/heads/{branch}.zip"
        
        logger.info("GitHub Indexer: downloading zip from %s", zip_url)
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, "repo.zip")

        try:
            # Download zip
            response = requests.get(zip_url, stream=True, timeout=30)
            if response.status_code == 404 and branch == "main":
                # Try fallback branch "master"
                zip_url = f"https://github.com/{owner}/{repo}/archive/refs/heads/master.zip"
                logger.info("GitHub Indexer: main branch failed, trying master branch from %s", zip_url)
                response = requests.get(zip_url, stream=True, timeout=30)

            if response.status_code != 200:
                raise RuntimeError(f"Failed to download repository zip from GitHub (HTTP {response.status_code})")

            with open(zip_path, "wb") as f:
                shutil.copyfileobj(response.raw, f)

            # Extract zip
            extract_dir = os.path.join(temp_dir, "extracted")
            with zipfile.ZipFile(zip_path, "r") as z:
                z.extractall(extract_dir)

            # Traverse and parse compatible files
            documents = []
            for root, _, files in os.walk(extract_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    rel_path = os.path.relpath(full_path, extract_dir)
                    # Skip binary assets/git files
                    if any(part.startswith(".") for part in rel_path.split(os.sep)):
                        continue
                    if file.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".ico", ".pdf", ".zip", ".tar", ".gz", ".mp3", ".wav", ".exe", ".dll", ".so", ".class", ".jar")):
                        # Skip large binary formats in git repo index
                        continue
                    
                    try:
                        doc = self.parser.parse_file(full_path)
                        # Decorate metadata with git path details
                        doc["metadata"]["source"] = f"github:{owner}/{repo}/{rel_path}"
                        doc["metadata"]["repo"] = f"{owner}/{repo}"
                        doc["metadata"]["relative_path"] = rel_path
                        documents.append(doc)
                    except Exception as e:
                        logger.warning("GitHub Indexer: skipped file '%s' due to error: %s", rel_path, e)
            
            logger.info("GitHub Indexer: parsed %d documents from repo '%s/%s'", len(documents), owner, repo)
            return documents
        finally:
            # Clean up temp files
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.error("GitHub Indexer: cleanup error: %s", e)
